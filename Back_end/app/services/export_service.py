import os
from app.core.config import settings
from minio import Minio
from io import BytesIO
from docx import Document
from app.core.celery_app import celery_app
import uuid

class ExportService:
    def __init__(self):
        self.client = None
        if settings.MINIO_ENDPOINT:
             try:
                self.client = Minio(
                    settings.MINIO_ENDPOINT.replace("http://", "").replace("https://", ""),
                    access_key=settings.MINIO_ACCESS_KEY,
                    secret_key=settings.MINIO_SECRET_KEY,
                    secure=False # Match your internal deployment
                )
             except Exception as e:
                 print(f"MinIO Initialization Error: {e}")

    def generate_markdown(self, data: dict) -> bytes:
        content = f"# Business Plan: {data.get('company_name', 'Untitled')}\n\n"
        for item in data.get('answers', []):
            content += f"## {item['question']}\n{item['answer']}\n\n"
        return content.encode("utf-8")

    def generate_docx(self, data: dict) -> bytes:
        doc = Document()
        doc.add_heading(f"Business Plan: {data.get('company_name', 'Untitled')}", 0)
        
        for item in data.get('answers', []):
            doc.add_heading(item['question'], level=1)
            doc.add_paragraph(item['answer'])
            
        file_stream = BytesIO()
        doc.save(file_stream)
        return file_stream.getvalue()

    def upload_to_storage(self, file_content: bytes, file_name: str, bucket: str = "exports"):
        if not self.client:
            return None
            
        try:
            # Ensure bucket exists
            if not self.client.bucket_exists(bucket):
                self.client.make_bucket(bucket)
            
            self.client.put_object(
                bucket,
                file_name,
                BytesIO(file_content),
                len(file_content)
            )
            return f"{bucket}/{file_name}"
        except Exception as e:
            print(f"Upload Error: {e}")
            return None

export_service = ExportService()

@celery_app.task(name="app.services.export_service.run_export_task")
def run_export_task(data: dict, format: str = "markdown"):
    """
    Background task for document generation and storage.
    """
    try:
        if format == "markdown":
            content = export_service.generate_markdown(data)
            ext = "md"
        elif format == "docx":
            content = export_service.generate_docx(data)
            ext = "docx"
        else:
            return {"status": "error", "message": f"Unsupported format: {format}"}
            
        file_name = f"export_{uuid.uuid4()}.{ext}"
        storage_path = export_service.upload_to_storage(content, file_name)
        
        return {
            "status": "success",
            "file_name": file_name,
            "storage_path": storage_path,
            "format": format
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}
